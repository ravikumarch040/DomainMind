from pathlib import Path

from docx import Document

from domainmind.data.models import RawDocument


def load_docx(path: Path) -> RawDocument:
    doc = Document(path)
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return RawDocument(
        text=text,
        source=str(path),
        metadata={"type": "docx"},
        doc_id=path.stem,
    )
